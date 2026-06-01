#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
2026年第5次部务会会议记录生成脚本 v5（最终版）
参照模板格式：
1. 出席/请假/列席人员只用姓名，同类用顿号隔开，不带职务
2. 议题用一、二、三……直接标出，内容不嵌套【（一）】格式
3. 列席人员标注在议题括号内
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_FZXBSJK = "方正小标宋简体"
FONT_HT = "黑体"
FONT_KT = "楷体"
FONT_FS = "仿宋"

def set_font(run, font_name, size, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_lnSpc(p, val=28):
    pPr = p._element.get_or_add_pPr()
    lSpc = OxmlElement('w:lnSpc')
    lSpcVal = OxmlElement('w:lnSpcVal')
    lSpcVal.set(qn('w:val'), str(val))
    lSpc.append(lSpcVal)
    pPr.append(lSpc)

def heading0(doc, text):
    """大标题 方正小标宋 22pt 居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_font(run, FONT_FZXBSJK, 22, bold=False)
    set_lnSpc(p)

def heading0sub(doc, text):
    """副标题 仿宋 16pt 居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, FONT_FS, 16)
    set_lnSpc(p)

def info(doc, text, sb=0, sa=3):
    """正文段落 仿宋 首行缩进"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run, FONT_FS, 16)
    set_lnSpc(p)
    return p

def info_noindent(doc, text, sb=0, sa=3):
    """正文段落 无首行缩进"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    run = p.add_run(text)
    set_font(run, FONT_FS, 16)
    set_lnSpc(p)
    return p

def agenda(doc, text):
    """议题标题 黑体 16pt 左对齐"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, FONT_HT, 16, bold=False)
    set_lnSpc(p)
    return p

def sub_agenda(doc, text):
    """分管汇报小标题 楷体 16pt"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, FONT_KT, 16, bold=False)
    set_lnSpc(p)
    return p

def item_num(doc, text, sb=0, sa=3):
    """编号条目 仿宋 首行缩进"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run, FONT_FS, 16)
    set_lnSpc(p)
    return p

def blank(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_lnSpc(p)

def create_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ===== 标题 =====
    heading0(doc, "中共五莲县委统战部部务会会议记录")
    heading0sub(doc, "（2026年第5次部务会）")

    # ===== 基本信息 =====
    info_noindent(doc, "时间：2026年5月27日（星期二）下午3:00", sa=3)
    info_noindent(doc, "地点：三楼会议室", sa=3)
    info_noindent(doc, "主持：张道伟", sa=3)

    info_noindent(doc, "出席人员：", sa=3)
    info_noindent(doc, "张道伟、徐军光、苑芳江、席光锋、徐慎文、李　兵、孙秀美", sb=0, sa=3)

    info_noindent(doc, "请假人员：", sa=3)
    info_noindent(doc, "潘姿安", sb=0, sa=3)

    info_noindent(doc, "列席人员：", sa=3)
    info_noindent(doc, "李国栋、朱芳、丰朔", sb=0, sa=12)

    # =========================================================================
    # 议题一
    # =========================================================================
    agenda(doc, "一、学习习近平总书记同美国总统特朗普、俄罗斯总统普京会谈精神，在中共中央政治局第二十五次集体学习时重要讲话精神")

    info(doc, "徐军光同志领学习近平总书记同美国总统特朗普会谈精神、在中央政治局第二十五次集体学习时的重要讲话精神。", sb=6, sa=6)
    info(doc, "会议指出，习近平总书记在同特朗普总统会谈时强调，中方致力于中美关系稳定、健康、可持续发展，赞同将构建中美建设性战略稳定关系作为中美关系新定位，为末来3年乃至更长时间的中美关系提供战略指引。中美建设性战略稳定关系应该是合作为主的积极稳定、竞争有度的良性稳定、分歧可控的常态稳定、和平可期的持久稳定。", sa=6)
    info(doc, "会议指出，习近平总书记在中央政治局第二十五次集体学习时强调，要站在统筹高质量发展和高水平安全的战略高度，充分认识做好防灾减灾救灾工作的重要性，着力提高防范应对自然灾害能力，切实维护人民群众生命财产安全。", sa=6)

    # =========================================================================
    # 议题二
    # =========================================================================
    agenda(doc, "二、学习《习近平总书记在浙江工作期间树立和践行正确政绩观的理论和实践》")

    info(doc, "苑芳江同志领学《习近平总书记在浙江工作期间树立和践行正确政绩观的理论和实践》。", sb=6, sa=6)
    info(doc, "会议指出，习近平总书记在浙江工作期间，始终大力倡导、带头践行正确政绩观，以高瞻远瞩的战略眼光、实事求是的科学态度、真抓实干的优良作风，引领推动浙江各项工作务实创新、全面发展。深入梳理总书记这一时期树立践行正确政绩观的理论和实践，主要体现在十一个方面：一是坚持立党为公、执政为民；二是坚持讲政治、顾大局；三是坚持调研开局、调研开路；四是坚持以战略眼光和系统思维谋发展、促发展；五是坚持把为人民谋利益作为树正风的根本目的；六是坚持从实际出发，按规律办事；七是坚持一张蓝图绘到底，一任接着一任干；八是坚持真抓实干，务求实效；九是坚持相信群众、依靠群众；十是坚持强化党性修养，加强党的建设；十一是坚持完善制度，建立科学考评体系。", sa=6)
    info(doc, "会议强调，要认真抓好防灾减灾救灾责任落实，各地区各部门要守土尽责，坚持统分结合、上下联动，推动形成集中共管、协同配合的工作格局。", sa=6)

    # =========================================================================
    # 议题三
    # =========================================================================
    agenda(doc, '三、审议《中共五莲县委统战部"三重一大"事项决策办法（试行）》（朱芳列席）')

    info(doc, "苑芳江同志汇报《中共五莲县委统战部'三重一大'事项决策办法（试行）》起草情况，对重大事项决策、重要干部任免、重大项目安排、大额资金使用等决策内容、决策程序、责任追究制度进行了详细说明。", sb=6, sa=6)
    info(doc, "李兵：同意", sa=4)
    info(doc, "徐军光：原则上同意本办法，第三页重要干部任免事项第四项，上级党代会代表都是党员，不涉及统战对象，应该是人大代表候选人、政协委员候选人，建议加上办公室共同负责。", sa=4)
    info(doc, "张道伟：我也同意，刚刚军光同志提的意见很好，建议修改完善后印发执行。", sa=6)

    # =========================================================================
    # 议题四
    # =========================================================================
    agenda(doc, "四、研究信息宣传工作（丰朔列席）")

    info(doc, "徐军光同志通报全县统战系统信息宣传工作有关情况。今年一季度，全市统战系统在省级以上媒体发布宣传稿件71条，日照市在全省位次处于第二；全县累计发稿120篇，在省以上媒体有突破，但横向对比其他区县还不突出。存在的主要问题：一是与媒体联动不紧密、宣传借力不够；二是常规动态类稿件多，总结类、专题类、典型人物类稿件少。", sb=6, sa=6)
    info(doc, "张道伟同志强调：", sa=4)
    info(doc, "一要吃透要求。认真研究上级对统战宣传信息工作的考核统计口径，连统计要求都吃不透，就没法做好对上宣传。要加强与宣传部、融媒体中心以及大众网的对接联系，学习借鉴他们的经验。", sa=4)
    info(doc, "二要成立专班。抽调两三个人，组成专门小组，成员包括大众网日照日报记者、融媒体中心记者、宣传部负责新闻的同志，加上统战部工作人员，专门研究统战宣传稿件的策划和推送。一星期开一次会，定题目、定任务、定媒体。", sa=4)
    info(doc, "三要借力发力。统战部长兼任宣传部长，要借这个优势，打好统战宣传主动仗。不能光靠自己写自己发，要让专业的人教我们写、帮我们写。要善于借力，打通人民网、新华社等中央级主流媒体渠道。", sa=4)
    info(doc, "四要挖掘亮点。统战工作的亮点是干出来的，也是总结出来的。要从民族团结进步创建、民营企业进边疆、太极拳大赛、寻美五莲等具体工作中提炼宣传亮点。要到乡镇、村居挖掘素材，发动基层宣传委员提供线索。第二季度宣传工作必须打个翻身仗。", sa=6)

    # =========================================================================
    # 议题五
    # =========================================================================
    agenda(doc, "五、分管领导工作汇报")

    sub_agenda(doc, "（一）李兵同志汇报民族宗教领域工作")
    item_num(doc, "1. 牧云谷石榴文化园项目：石榴树全部栽植完成，因雨水充足，成活率良好，部分已开花，正在推进后续配套工作。", sb=4, sa=3)
    item_num(doc, "2. 国家民委旅游促进各民族交往交流交融实践项目：依托大青山太极拳大赛，整合资源，争取申报成功。", sa=3)
    item_num(doc, "3. 民品贴息贷款绩效评价：省财政厅、省民委委托第三方对五征集团2023至2025年度民品贴息贷款开展绩效评价，2023、2024、2025年度拨款已全部清零，评价工作正在进行。", sa=3)
    item_num(doc, "4. 对口支援麦盖提县工作：已召开座谈会，梳理可对接的教育、卫生、产业等合作项目。", sa=3)
    item_num(doc, "5. 非法宗教活动查处：高效处理两起非法宗教活动。", sa=3)
    item_num(doc, "6. 党员县级联审：配合组织部、宣传部、教体局完成近500名党员县级联审。", sa=3)
    item_num(doc, "7. 省委巡视反馈问题整改：已通过验收。", sa=3)
    item_num(doc, "8. 光明寺土地使用协议：参照外地经验，协调封管委与光明寺签订土地使用协议，解决宗教场所无房产证问题。", sa=3)
    item_num(doc, "9. 宗教活动场所全覆盖调研：已初步完成，发现的机构建设、人员管理、财务管理、安全管理等问题均已整改完成。", sa=6)

    sub_agenda(doc, "（二）徐慎文同志汇报侨务和新的社会阶层人士统战工作")
    item_num(doc, "1. 侨联工作：与县侨联邓主席、莫主席对接，推进侨务之家和基层侨联组织建设；对接通州区教育局开展圆梦行动；陪同亚太森博负责人到五莲一中、五莲中学对接；填报侨助力双招双引工作情况；在校华人华侨学生信息摸底；落实大青山太极讲座活动。", sb=4, sa=3)
    item_num(doc, "2. 侨情数据调研：侨联系统数据调研工作部署座谈交流会召开，填报时间至10月底，中国侨联要求详细填报，正在梳理核实。", sa=3)
    item_num(doc, "3. 新联会工作：完成新联会年审；推进乡镇新联会备案或注册工作；召开新联会会长会议；落实市新联会理事会会议精神；推荐市网联会会员；对接维光助学公益捐赠（计划捐建吉他教室和无人机教室）。", sa=3)
    item_num(doc, "4. 市网联会筹建：配合做好市网联会会员推荐工作。", sa=6)

    sub_agenda(doc, "（三）席光锋同志汇报分管领域工作")
    item_num(doc, "1. 国企成长工作会议：组织城发集团等三家国有企业召开会议，部署国企成长相关工作。", sb=4, sa=3)
    item_num(doc, "2. 质量工作会议：与城发集团对接，指导其完善质量会办公室、制度版面等准备工作。", sa=3)
    item_num(doc, '3. 民主党派工作：筹备召开全县民主党派"参政为公、实干为民"主题教育座谈会。', sa=3)
    item_num(doc, "4. 党外人士谈话材料：上报市委统战部领导干部与党外人士谈话材料。", sa=3)
    item_num(doc, "5. 全市社院院长座谈会：准备发言材料。", sa=3)
    item_num(doc, "6. 白鹭湾美术馆教育基地：上报省级教育基地教学点开展培训情况。", sa=3)
    item_num(doc, "7. 黄炎培职业教育创新创业大赛：联合教体局、两个职业院校召开推进会议，11篇案例上报市里。", sa=3)
    item_num(doc, "8. 科技中专产教融合：新华社、学习强国等媒体报道。", sa=3)
    item_num(doc, "9. 下步工作：跟进黄炎培大赛闭幕式；推进五征集团党委政治类联谊会工作；考虑职业院校换届事宜。", sa=6)

    sub_agenda(doc, "（四）苑芳江同志汇报分管领域工作")
    item_num(doc, "1. 统战工作责任制自查自纠：已向各党委党组下发通知和督查提纲，6月20日前完成自查自纠，6月10日后省市下来抽查。", sb=4, sa=3)
    item_num(doc, "2. 巡查工作配合：已完成机关党委先进基层党组织材料上报、青年干部上讲台（李国栋讲机关公文格式）、巡查协助配合等工作。", sa=3)
    item_num(doc, "3. 茶文化研习周活动：5月29日接待参访团到刘罗湾开展活动。", sa=3)
    item_num(doc, "4. 老干部工作：郑德军同志7月份退休，正在准备退休申请。", sa=6)

    # =========================================================================
    # 议题六
    # =========================================================================
    agenda(doc, "六、张道伟同志讲话")

    info(doc, "一要深入学习。要把学习习近平总书记关于做好新时代党的统一战线工作的重要论述作为重要政治任务，深刻领会其核心要义和实践要求，切实用以武装头脑、指导实践、推动工作。", sb=6, sa=6)
    info(doc, "二要突出重点、打造亮点。要聚焦民族团结进步创建、宗教领域和谐稳定、新的社会阶层人士统战等重点工作，集中力量、精准发力，力争在重点领域取得突破性进展，形成一批有影响力的工作品牌。", sa=6)
    info(doc, "三要强化自身建设。要持续加强理论学习，提升业务能力，严守纪律规矩，打造政治坚定、业务精通、作风过硬的统战干部队伍。要加强与省市统战部门的沟通对接，积极争取支持，推动全县统战工作再上新台阶。", sa=6)
    info(doc, '四要抓好工作落实。对本次会议研究的各项议题，要明确责任分工和完成时限，定期调度通报，确保各项工作任务落地见效。各分管领导要认真履行"一岗双责"，既抓业务又抓党建，以钉钉子精神推动各项任务落实。', sa=6)
    info(doc, "五要守好安全底线。始终绷紧安全这根弦，毫不松懈抓好宗教领域安全、生产安全、意识形态安全等各项工作，确保全县统一战线领域和谐稳定。", sa=12)

    # ===== 结尾 =====
    blank(doc)
    info_noindent(doc, "（以下无内容）", sb=0, sa=6)
    blank(doc)
    info_noindent(doc, "中共五莲县委统战部", sb=12, sa=3)
    info_noindent(doc, "2026年5月27日", sb=0, sa=0)

    output_path = "/mnt/nfs/2026年统战工作/1.办公室/9.部务会/5.27/会议记录_2026年第5次部务会.docx"
    doc.save(output_path)
    print(f"最终版已保存：{output_path}")

create_doc()
