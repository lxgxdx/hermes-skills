from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

section = doc.sections[0]
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

def set_font(run, font='仿宋_GB2312', size=16, bold=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)

def body_para(doc, text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else 0
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run)
    return p

def blank(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(28)

def title_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.bold = False
    run.font.name = '方正小标宋简体'
    run.font.size = Pt(22)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')

def heading1(doc, text):
    """一级标题：黑体不加粗"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.bold = False
    run.font.name = '黑体'
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def heading2(doc, text):
    """二级标题：楷体不加粗"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.bold = False
    run.font.name = '楷体_GB2312'
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')

def heading3(doc, text):
    """三级标题：仿宋加粗"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.bold = True
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

def info_line(doc, label, content):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label)
    set_font(r1, bold=True)
    r2 = p.add_run(content)
    set_font(r2)

def attendee(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run)

def bullet(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run)
