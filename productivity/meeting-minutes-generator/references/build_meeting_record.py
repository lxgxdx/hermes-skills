#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
会议记录生成脚本模板 — 2026年第5次部务会版（v4，已验证正确）
用法：复制到工作目录，修改内容后运行 python3 build_meeting_record.py

关键修正（踩坑记录）：
1. 议题表 .doc 格式必须用 libreoffice 转 txt，不能 python-docx 直接读
2. 文字稿说话人编号 ≠ 议题表领学顺序，必须交叉验证
3. 末位发言顺序：潘姿安→李兵→徐军光→张道伟（本次潘姿安请假，实际：李兵→徐军光→张道伟）
4. 表决记录：同意+原话，无原话只写同意
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_FZXBSJK = "方正小标宋简体"
FONT_HT = "黑体"
FONT_KT = "楷体"
FONT_FS = "仿宋"

def set_font(run, font_name, size, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_lnSpc(p, val=28):
    pPr = p._element.get_or_add_pPr()
    lSpc = OxmlElement('w:lnSpc')
    lSpcVal = OxmlElement('w:lnSpcVal')
    lSpcVal.set(qn('w:val'), str(val))
    lSpc.append(lSpcVal)
    pPr.append(lSpc)

def h0(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_font(run, FONT_FZXBSJK, 22)
    set_lnSpc(p)

def h0sub(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, FONT_FS, 16)
    set_lnSpc(p)

def info(doc, text, indent=True, sb=0, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run, FONT_FS, 16)
    set_lnSpc(p)

def h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, FONT_HT, 16)
    set_lnSpc(p)

def h2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("【" + text + "】")
    set_font(run, FONT_KT, 16)
    set_lnSpc(p)

def h3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("（" + text + "）")
    set_font(run, FONT_FS, 16, bold=True)
    set_lnSpc(p)

def blank(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_lnSpc(p)

def create_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # 标题
    h0(doc, "中共五莲县委统战部部务会会议记录")
    h0sub(doc, "（2026年第5次部务会）")

    # 基本信息
    info(doc, "时间：2026年5月27日（星期二）下午3:00", indent=False, sa=3)
    info(doc, "地点：三楼会议室", indent=False, sa=3)
    info(doc, "主持：张道伟（县委常委、统战部部长）", indent=False, sa=3)

    # 出席人员
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("出席人员：")
    set_font(run, FONT_FS, 16, bold=True)
    set_lnSpc(p)

    members = [
        ("张道伟", "县委常委、统战部部长"),
        ("徐军光", "副部长，主持日常工作"),
        ("苑芳江", "副部长，分管党建、共青团妇、精神文明"),
        ("席光锋", "副部长，分管党外知识分子、民主党派、职教社"),
        ("徐慎文", "副部长，分管新的社会阶层、侨务、侨联"),
        ("李　兵", "副部长，分管民族宗教"),
    ]
    for name, desc in members:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.74)
        run = p.add_run(f"{name}　{desc}")
        set_font(run, FONT_FS, 16)
        set_lnSpc(p)

    # 请假
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("请假人员：")
    set_font(run, FONT_FS, 16, bold=True)
    set_lnSpc(p)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run("潘姿安　副部长，分管办公室、台港澳、信息宣传")
    set_font(run, FONT_FS, 16)
    set_lnSpc(p)

    # 列席
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("列席人员：")
    set_font(run, FONT_FS, 16, bold=True)
    set_lnSpc(p)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run("李国栋　办公室主任（不发言，负责记录）")
    set_font(run, FONT_FS, 16)
    set_lnSpc(p)

    info(doc, "记录人：孙秀美（不发言，专职记录）", indent=False, sb=3, sa=12)

    # ===== 议题一 =====
    h1(doc, "议题一：学习习近平总书记同美国总统特朗普、俄罗斯总统普京会谈精神，在中共中央政治局第二十五次集体学习时重要讲话精神")
    h2(doc, "（一）领学")
    info(doc, "徐军光同志领学习近平总书记同美国总统特朗普会谈精神、在中央政治局第二十五次集体学习时的重要讲话精神。", sb=6, sa=6)
    h2(doc, "（二）会议强调")
    info(doc, "会议指出，习近平总书记在同特朗普总统会谈时强调，中方致力于中美关系稳定、健康、可持续发展，赞同将构建中美建设性战略稳定关系作为中美关系新定位。", sb=6, sa=6)
    info(doc, "会议指出，习近平总书记在中央政治局第二十五次集体学习时强调，要站在统筹高质量发展和高水平安全的战略高度，充分认识做好防灾减灾救灾工作的重要性。", sa=6)

    # ===== 议题二 =====
    h1(doc, "议题二：学习《习近平总书记在浙江工作期间树立和践行正确政绩观的理论和实践》")
    h2(doc, "（一）领学")
    info(doc, "苑芳江同志领学《习近平总书记在浙江工作期间树立和践行正确政绩观的理论和实践》。", sb=6, sa=6)
    h2(doc, "（二）会议强调")
    info(doc, "会议强调，习近平总书记在浙江工作期间树立践行正确政绩观主要体现在十一个方面：一是坚持立党为公、执政为民；二是坚持讲政治、顾大局；三是坚持调研开局、调研开路；四是坚持以战略眼光和系统思维谋发展；五是坚持把为人民谋利益作为树正风的根本目的；六是坚持从实际出发，按规律办事；七是坚持一张蓝图绘到底；八是坚持真抓实干，务求实效；九是坚持相信群众、依靠群众；十是坚持强化党性修养；十一是坚持完善制度，建立科学考评体系。", sb=6, sa=6)
    info(doc, "会议强调，要认真抓好防灾减灾救灾责任落实，守土尽责，坚持统分结合、上下联动，推动形成集中共管、协同配合的工作格局。", sa=6)

    # ===== 议题三 =====
    h1(doc, "议题三：审议《中共五莲县委统战部\"三重一大事项\"决策办法（试行）》")
    h2(doc, "（一）汇报情况")
    info(doc, "苑芳江同志汇报《中共五莲县委统战部\"三重一大事项\"决策办法（试行）》起草情况。朱芳同志列席。", sb=6, sa=6)
    h2(doc, "（二）审议意见")
    info(doc, "徐慎文同志：第三页重要干部任免事项第四项，上级党代会代表都是党员，不涉及统战对象，建议加上办公室共同负责。", sb=6, sa=4)
    info(doc, "张道伟同志：我们要参照执行县委新发的三重一大决策办法。决策征求意见条款要根据公众影响范围和程度区分对待，统战部有的事项涉及面广需要征求，有的一般工作不需要征求，要研究清楚再定。", sa=4)
    info(doc, "徐军光同志：决策程序跟县里保持一致，征求意见应区分保密要求。", sa=4)
    info(doc, "张道伟同志：修改完善后印发执行。既然有这个办法了，该上会研究的就要上会研究。", sa=4)
    h2(doc, "（三）表决结果")
    info(doc, "经审议，原则通过，修改完善后印发执行。", sb=6, sa=12)

    # ===== 议题四 =====
    h1(doc, "议题四：研究信息宣传工作")
    h2(doc, "（一）情况通报")
    info(doc, "徐军光同志通报全县统战系统信息宣传工作有关情况。一季度全市统战系统在省级以上媒体发布宣传稿件71条，日照市在全省位次处于第二；全县累计发稿120篇，存在的主要问题：一是与媒体联动不紧密；二是常规动态类稿件多，总结类、专题类稿件少。丰朔同志列席。", sb=6, sa=6)
    h2(doc, "（二）领导要求")
    info(doc, "张道伟同志强调：", sb=6, sa=4)
    info(doc, "一要吃透要求。认真研究上级对统战宣传工作的考核统计口径，加强与宣传部、融媒体中心以及大众网的对接联系。", sa=4)
    info(doc, "二要成立专班。抽调两三个人，组成专门小组，一星期开一次会，定题目、定任务、定媒体。", sa=4)
    info(doc, "三要借力发力。借统战部长兼任宣传部长这个优势，打通人民网、新华社等中央级主流媒体渠道。", sa=4)
    info(doc, "四要挖掘亮点。从民族团结进步创建、民营企业进边疆、太极拳大赛、寻美五莲等具体工作中提炼宣传亮点，第二季度宣传工作必须打个翻身仗。", sa=6)

    # ===== 议题五 =====
    h1(doc, "议题五：分管领导工作汇报")
    h2(doc, "（一）李兵同志汇报民族宗教领域工作")
    info(doc, "李兵同志汇报民族宗教领域近期工作（共九项）：", sb=6, sa=4)
    info(doc, "1. 穆园古石榴文化园项目：石榴树全部栽植完成，成活率良好。", sa=3)
    info(doc, "2. 国家民委旅游促进各民族交往交流交融实践项目：依托大青山太极拳大赛整合资源申报。", sa=3)
    info(doc, "3. 民品贴息贷款绩效评价：省第三方已启动五征集团2023至2025年度绩效评价，拨款已全部清零。", sa=3)
    info(doc, "4. 对口支援麦盖提县工作：已召开座谈会，梳理可对接的合作项目。", sa=3)
    info(doc, "5. 非法宗教活动查处：高效处理两起非法宗教活动，配合公安部门依法处理。", sa=3)
    info(doc, "6. 党员县级联审：配合完成近500名党员联审。", sa=3)
    info(doc, "7. 省委巡视反馈问题整改：已通过验收。", sa=3)
    info(doc, "8. 光明寺土地使用协议：协调封管委与光明寺签订协议，解决场所无房产证问题。", sa=3)
    info(doc, "9. 宗教活动场所全覆盖调研：已完成，发现问题均已整改。", sa=6)

    h2(doc, "（二）徐慎文同志汇报侨务和新的社会阶层人士统战工作")
    info(doc, "徐慎文同志汇报侨务和新联会工作：", sb=6, sa=4)
    info(doc, "1. 侨联工作：推进侨务之家和基层侨联组织建设，开展圆梦行动，对接大青山太极讲座活动。", sa=3)
    info(doc, "2. 侨情数据调研：中国侨联要求详细填报，正在梳理核实，填报至10月底。", sa=3)
    info(doc, "3. 新联会工作：完成年审，推进乡镇新联会备案，推荐市网联会会员，落实维光助学公益捐赠。", sa=3)
    info(doc, "4. 市网联会筹建：配合做好会员推荐工作。", sa=6)

    h2(doc, "（三）席光锋同志汇报分管领域工作")
    info(doc, "席光锋同志汇报党建、干部人事、民主党派和职教社工作：", sb=6, sa=4)
    info(doc, "1. 国企成长工作会议：组织城发集团等三家国有企业部署相关工作。", sa=3)
    info(doc, "2. 民主党派工作：筹备全县民主党派\"参政为公、实干为民\"主题教育座谈会。", sa=3)
    info(doc, "3. 黄炎培职业教育创新创业大赛：联合教体局、职业院校召开推进会议，11篇案例上报市里。", sa=3)
    info(doc, "4. 下步工作：跟进黄炎培大赛闭幕式，推进五征集团党委政治类联谊会工作。", sa=6)

    h2(doc, "（四）苑芳江同志汇报分管领域工作")
    info(doc, "苑芳江同志汇报党建、共青团妇、精神文明工作：", sb=6, sa=4)
    info(doc, "1. 统战工作责任制自查自纠：已下发通知，6月20日前完成自查，6月10日后省市抽查。", sa=3)
    info(doc, "2. 巡查配合工作：完成先进基层党组织材料上报，组织青年干部上讲台。", sa=3)
    info(doc, "3. 茶文化研习周活动：5月29日接待参访团到刘罗湾开展活动。", sa=6)

    h2(doc, "（五）主要领导交办事项")
    info(doc, "张道伟同志交办：太极拳大赛融入统战元素；新联会、侨联每年招引规上项目。", sb=6, sa=12)

    # ===== 议题六 =====
    h1(doc, "议题六：主要领导讲话")
    h2(doc, "张道伟同志总结讲话")
    info(doc, "一要深入学习习近平总书记关于做好新时代党的统一战线工作的重要论述，切实用以武装头脑、指导实践、推动工作。", sb=6, sa=6)
    info(doc, "二要突出重点、打造亮点，聚焦民族团结进步创建、宗教领域和谐稳定、新的社会阶层人士统战等重点工作，形成一批有影响力的工作品牌。", sa=6)
    info(doc, "三要强化自身建设，加强与省市统战部门的沟通对接，推动全县统战工作再上新台阶。", sa=6)
    info(doc, "四要抓好工作落实，各分管领导要认真履行\"一岗双责\"，以钉钉子精神推动各项任务落实。", sa=6)
    info(doc, "五要守好安全底线，确保全县统一战线领域和谐稳定。", sa=12)

    # 结尾
    blank(doc)
    info(doc, "（以下无内容）", indent=False, sb=0, sa=6)
    blank(doc)
    info(doc, "中共五莲县委统战部", indent=False, sb=12, sa=3)
    info(doc, "2026年5月27日", indent=False, sb=0, sa=0)

    output_path = "/mnt/nfs/2026年统战工作/1.办公室/9.部务会/5.27/会议记录_2026年第5次部务会_v4.docx"
    doc.save(output_path)
    print(f"已保存：{output_path}")

create_doc()
