#!/usr/bin/env python3
"""
高级文档编辑工具 - 深度调整Word和Excel格式
支持公文格式排版、表格美化等
符合机关公文格式规范国家标准
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Border, Side, PatternFill, Alignment
from openpyxl.utils import get_column_letter
import argparse

# ============================================
# 公文字体配置（机关公文格式规范国家标准）
# ============================================
GONGWEN_FONTS = {
    "一级标题": {"name": "黑体", "size": Pt(15), "bold": True},       # 一、xxx → 黑体三号
    "二级标题": {"name": "楷体_GB2312", "size": Pt(15), "bold": True}, # （一）xxx → 楷体三号
    "三级标题": {"name": "仿宋_GB2312", "size": Pt(15), "bold": True}, # 三是xxx → 仿宋三号加粗
    "正文": {"name": "仿宋_GB2312", "size": Pt(15), "bold": False},   # 正文 → 仿宋三号
    "页码": {"name": "宋体", "size": Pt(12), "bold": False},          # 页码 → 四号宋体
}

# 公文页面设置（国家标准）
GONGWEN_PAGE = {
    "top": Cm(3.7),      # 上边距 3.7cm
    "bottom": Cm(3.5),   # 下边距 3.5cm
    "left": Cm(2.8),     # 左边距 2.8cm
    "right": Cm(2.6),    # 右边距 2.6cm
}

# 公文段落格式
GONGWEN_PARAGRAPH = {
    "line_spacing": Pt(28),      # 固定值28磅
    "first_line_indent": Cm(0.74),  # 首行缩进2个字符（仿宋三号约0.74cm）
    "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,  # 两端对齐
}


class WordDocumentEditor:
    """Word文档编辑器 - 符合公文格式规范"""
    
    def __init__(self, filepath=None):
        self.doc = Document() if filepath is None else Document(filepath)
        # 默认设置公文页面
        self.set_page_setup()
    
    def set_page_setup(self, top=3.7, bottom=3.5, left=2.8, right=2.6, orientation="portrait"):
        """设置公文页面布局（国家标准）"""
        section = self.doc.sections[0]
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)
        
        if orientation == "landscape":
            section.page_width, section.page_height = section.page_height, section.page_width
        
        return self
    
    def set_title(self, title, level=1):
        """
        设置标题
        level=1: 一级标题（黑体三号），如"一、放假安排"
        level=2: 二级标题（楷体三号），如"（一）值班安排"
        level=3: 三级标题（仿宋三号加粗），如"一是安全检查"
        """
        p = self.doc.add_paragraph()
        
        # 根据级别设置前缀和字体
        if level == 1:
            title_text = f"一、{title}"
            font_config = GONGWEN_FONTS["一级标题"]  # 黑体三号
        elif level == 2:
            title_text = f"（一）{title}"
            font_config = GONGWEN_FONTS["二级标题"]  # 楷体三号
        elif level == 3:
            title_text = f"一是{title}"
            font_config = GONGWEN_FONTS["三级标题"]  # 仿宋三号加粗
        else:
            title_text = title
            font_config = GONGWEN_FONTS["正文"]
        
        run = p.add_run(title_text)
        run.font.name = font_config["name"]
        run.font.size = font_config["size"]
        run.font.bold = font_config["bold"]
        
        # 标题后空一行
        p.paragraph_format.space_after = Pt(16)
        
        return p
    
    def add_heading(self, text, level=0):
        """
        添加正文标题（不带数字标记）
        level=0: 文档总标题 - 方正小标宋二号居中
        level=1-9: 各级标题（不带前缀，由set_title添加）
        """
        p = self.doc.add_heading(text, level=level)
        
        if level == 0:
            # 文档总标题：方正小标宋二号，居中
            # 注意：需要安装"方正小标宋"字体，系统默认使用Noto Serif CJK SC Bold替代
            run = p.runs[0] if p.runs else p.add_run(text)
            run.font.name = "Noto Serif CJK SC Bold"  # 优先使用Noto Serif CJK替代
            run.font.size = Pt(22)
            run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(24)
        
        return p
    
    def add_paragraph(self, text, font_name=None, font_size=None, 
                     align=None, bold=False, first_line_indent=None,
                     line_spacing=None, space_before=0, space_after=0):
        """
        添加正文段落
        
        默认格式（公文标准）：
        - 字体：仿宋_GB2312
        - 字号：三号（15磅）
        - 对齐：两端对齐
        - 首行缩进：2字符
        - 行距：固定28磅
        """
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        
        # 字体设置
        if font_name is None:
            font_name = GONGWEN_FONTS["正文"]["name"]
        if font_size is None:
            font_size = GONGWEN_FONTS["正文"]["size"]
        
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        
        # 对齐方式
        if align is None:
            align = GONGWEN_PARAGRAPH["alignment"]
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 首行缩进
        if first_line_indent is None:
            first_line_indent = True
        if first_line_indent:
            p.paragraph_format.first_line_indent = GONGWEN_PARAGRAPH["first_line_indent"]
        
        # 行距
        if line_spacing is None:
            line_spacing = GONGWEN_PARAGRAPH["line_spacing"]
        p.paragraph_format.line_spacing = line_spacing
        
        # 段前段后间距
        if space_before > 0:
            p.paragraph_format.space_before = Pt(space_before)
        if space_after > 0:
            p.paragraph_format.space_after = Pt(space_after)
        
        return p
    
    def add_empty_line(self, count=1):
        """添加空行"""
        for _ in range(count):
            self.add_paragraph(" ", space_after=0)
    
    def add_numbered_list(self, items, level=1):
        """
        添加编号列表
        level=1: 一级编号（一、二、三、）
        level=2: 二级编号（（一）（二）（三））
        level=3: 三级编号（1. 2. 3.）
        """
        for i, item in enumerate(items, 1):
            if level == 1:
                prefix = f"{item}、"
            elif level == 2:
                prefix = f"（{item}）"
            elif level == 3:
                prefix = f"{item}."
            else:
                prefix = f"{item}、"
            
            self.add_paragraph(f"{prefix}{item}", first_line_indent=True)
    
    def add_table(self, data, header_row=True, style="Table Grid"):
        """添加表格"""
        table = self.doc.add_table(rows=1, cols=len(data[0]))
        table.style = style
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        for i, key in enumerate(data[0]):
            hdr_cells[i].text = str(key)
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "黑体"
                    run.font.size = Pt(14)
                    run.font.bold = True
        
        # 填充数据
        for row_idx, row_data in enumerate(data[1:], start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = str(cell_data)
        
        return table
    
    def add_attachment_note(self, attachments):
        """
        添加附件说明
        attachments: 附件列表，如 ["1. XXX文件", "2. YYY表格"]
        """
        self.add_paragraph("附件：")
        for i, attachment in enumerate(attachments, 1):
            self.add_paragraph(f"  {attachment}", first_line_indent=False)
        return self
    
    def add_sender_info(self, sender, date, align="right"):
        """
        添加发文机关署名和成文日期
        
        Args:
            sender: 发文机关名称
            date: 成文日期，如 "2026年2月1日"
            align: 对齐方式（right/left）
        """
        self.add_empty_line()
        
        # 成文日期（右侧）
        if align == "right":
            self.add_paragraph(date, align="right", first_line_indent=False, space_before=12)
        else:
            self.add_paragraph(date, first_line_indent=False, space_before=12)
        
        # 发文机关署名（日期上面）
        if align == "right":
            self.add_paragraph(sender, align="right", first_line_indent=False)
        else:
            self.add_paragraph(sender, first_line_indent=False)
        
        return self
    
    def add_page_number(self, format_str="—　{page}　—"):
        """
        添加页码
        格式：—　1　—（四号宋体）
        注意：首页、封面、目录不标页码
        """
        # 简化实现：添加页码占位符
        # 实际排版中需在页脚中设置
        pass
    
    def save(self, filepath):
        """保存文档"""
        self.doc.save(filepath)
        print(f"✅ Word文档已保存: {filepath}")


class ExcelDocumentEditor:
    """Excel文档编辑器"""
    
    def __init__(self, filepath=None):
        self.wb = Workbook() if filepath is None else load_workbook(filepath)
        self.ws = self.wb.active
    
    def add_data(self, data, start_row=1, start_col=1):
        """添加数据"""
        for row_idx, row_data in enumerate(data, start=start_row):
            for col_idx, cell_value in enumerate(row_data, start=start_col):
                cell = self.ws.cell(row=row_idx, column=col_idx, value=cell_value)
        return self
    
    def format_header(self, row=1, font_name="宋体", font_size=11, 
                      bold=True, bg_color="D9D9D9", align="center"):
        """格式化表头"""
        header_fill = PatternFill(start_color=bg_color, end_color=bg_color, fill_type="solid")
        
        for cell in self.ws[row]:
            cell.font = Font(name=font_name, size=font_size, bold=bold)
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal=align, vertical="center")
        
        return self
    
    def set_column_width(self, col, width):
        """设置列宽"""
        self.ws.column_dimensions[get_column_letter(col)].width = width
        return self
    
    def set_row_height(self, row, height):
        """设置行高"""
        self.ws.row_dimensions[row].height = height
        return self
    
    def add_borders(self, border_type="all"):
        """添加边框"""
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for row in self.ws.iter_rows(min_row=1, max_row=self.ws.max_row, 
                                   min_col=1, max_col=self.ws.max_column):
            for cell in row:
                if border_type == "all":
                    cell.border = thin_border
                elif border_type == "outside":
                    pass  # 简化处理
        return self
    
    def format_numbers(self, col, format_string="#,##0.00"):
        """格式化数字"""
        for cell in self.ws[get_column_letter(col)][1:]:
            if isinstance(cell.value, (int, float)):
                cell.number_format = format_string
        return self
    
    def set_cell_font(self, row, col, font_name="宋体", font_size=11, 
                      color=None, bold=False):
        """设置单元格字体"""
        cell = self.ws.cell(row=row, column=col)
        cell.font = Font(name=font_name, size=font_size, bold=bold)
        if color:
            cell.font.color = RGBColor(*color)
        return self
    
    def save(self, filepath):
        """保存文档"""
        self.wb.save(filepath)
        print(f"✅ Excel文档已保存: {filepath}")


def main():
    parser = argparse.ArgumentParser(description="高级文档编辑工具（公文格式规范）")
    subparsers = parser.add_subparsers(dest="command", help="子命令")
    
    # Word命令
    word_parser = subparsers.add_parser("word", help="Word文档编辑")
    word_parser.add_argument("--input", "-i", help="输入文件")
    word_parser.add_argument("--output", "-o", help="输出文件")
    word_parser.add_argument("--title", "-t", help="设置标题")
    word_parser.add_argument("--content", "-c", help="添加内容")
    word_parser.add_argument("--font", "-f", default="仿宋_GB2312", help="字体")
    word_parser.add_argument("--size", "-s", type=int, default=15, help="字号（三号为15）")
    
    # Excel命令
    excel_parser = subparsers.add_parser("excel", help="Excel文档编辑")
    excel_parser.add_argument("--input", "-i", help="输入文件")
    excel_parser.add_argument("--output", "-o", help="输出文件")
    excel_parser.add_argument("--data", "-d", help="CSV格式数据")
    excel_parser.add_argument("--header-row", action="store_true", help="有表头")
    
    args = parser.parse_args()
    
    if args.command == "word":
        editor = WordDocumentEditor(args.input)
        if args.title:
            editor.set_title(args.title)
        if args.content:
            editor.add_paragraph(args.content, font_name=args.font, font_size=Pt(args.size))
        output = args.output or (args.input.replace(".docx", "_edited.docx") if args.input else "output.docx")
        editor.save(output)
    
    elif args.command == "excel":
        editor = ExcelDocumentEditor(args.input)
        if args.data:
            data = [row.split(",") for row in args.data.split(";")]
            editor.add_data(data)
        output = args.output or (args.input.replace(".xlsx", "_edited.xlsx") if args.input else "output.xlsx")
        editor.save(output)
    
    else:
        print("📝 使用说明:")
        print("  python3 editor.py word -t '标题' -c '内容' -o output.docx")
        print("  python3 editor.py excel -d '姓名,年龄,城市' -o output.xlsx")


if __name__ == "__main__":
    main()
